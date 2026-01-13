from docx import Document
from pathlib import Path

def save_to_word(content, output_path):
    """Creates a Word document and writes the given content into it."""
    try:
        document = Document()
        document.add_heading("Generated Content", level=1)
        document.add_paragraph(content)
        
        # Ensure the output directory exists
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        document.save(output_path)
        print(f"Word document saved to: {output_path}")
    except Exception as e:
        print(f"An error occurred while saving the Word document: {e}")