import os
from openai import AzureOpenAI
from pathlib import Path
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document

load_dotenv()
# --- Configuration ---
# Ensure the environment variables are set (see Prerequisites above)
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
DEPLOYMENT_NAME = os.getenv("DEPLOYMENT_NAME") # Replace with your model deployment name (e.g., "gpt-35-turbo")

if not all([AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_KEY, DEPLOYMENT_NAME]):
    print("Error: Required environment variables or deployment name not set.")
    exit()

# --- Initialize the Azure OpenAI Client ---
client = AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version="2025-04-01-preview" # Use a compatible API version
)

def resume_agent():
    """Reads a file and sends it to Azure AI for initial analysis."""
    rif = Path.cwd()/"input_files/RIF.pdf"
    rif_data = extract_form_data(rif)
    notes = Path.cwd()/"input_files/notes.txt"
    notes_content = notes.read_text(encoding="utf-8") # Standard in 2026
    try:
        # Load the file content
        transcript_path = Path.cwd()/"input_files/transcript.txt"
        transcript_content = transcript_path.read_text(encoding="utf-8") # Standard in 2026
        
        # Create a specialized prompt for analysis
        analysis_request = f"Please create a resume in an optimized Microsoft Word format from the provided information form, transcript, and interview notes. The summary should include 1-3 detailed bullet points for each of the following sections: Skills, Experience, Education, and Projects (if applicable):\n\n{rif_data}\n\nTranscript:\n\n{transcript_content}\n\nManager's notes:\n{notes_content}"
        conversation_history = []
        # Add the file content to history as a user message
        conversation_history.append({"role": "user", "content": analysis_request})
        
        # Call Azure OpenAI API
        response = client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=conversation_history,
            max_completion_tokens=30000, # Increased for analysis
        )
        
        assistant_message = response.choices[0].message.content
        conversation_history.append({"role": "assistant", "content": assistant_message})
        mk = conversation_history[-1]["content"]        
        # Save the assistant message to a Word document
        output_path = Path.cwd() / "resume_files/resume.docx"
        save_to_word(mk, output_path)
    except FileNotFoundError:
        print(f"Error: The file '{transcript_path}' was not found.")
    except Exception as e:
        print(f"An error occurred: {e}")
        
    print("Resume generation complete!")

def extract_form_data(pdf_path):
    with open(pdf_path, "rb") as file:
        reader = PdfReader(file)
        # get_fields() returns a dictionary of form fields and their values
        form_fields = reader.get_fields()
        return form_fields

def save_to_word(content, output_path):
    """Creates a Word document and writes the given content into it."""
    try:
        document = Document()
        document.add_heading("Generated Content", level=1)
        document.add_paragraph(content)
        
        # Ensure the output directory exists
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        document.save(output_path)
        print(f"Word document saved to: {output_path}")
    except Exception as e:
        print(f"An error occurred while saving the Word document: {e}")
        
# --- Main Analysis Loop ---
if __name__ == "__main__":
    resume_agent()